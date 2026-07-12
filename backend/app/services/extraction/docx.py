"""DOCX extractor backed by python-docx.

Word documents have no reliable page boundaries (pagination is a rendering-time
concern), so we emit a single logical page. Paragraphs and table cell text are
flattened in reading order.
"""

from __future__ import annotations

from pathlib import Path

from app.core.exceptions import ExtractionError
from app.models.document import Document
from app.services.extraction.base import BaseExtractor


class DocxExtractor(BaseExtractor):
    extensions = (".docx",)
    mime_types = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    def extract(self, path: Path) -> Document:
        try:
            import docx  # python-docx
        except ImportError as exc:  # pragma: no cover - env dependent
            raise ExtractionError(
                "python-docx is required for DOCX extraction. Install with "
                "`pip install python-docx`."
            ) from exc

        try:
            document = docx.Document(str(path))
        except Exception as exc:
            raise ExtractionError(f"failed to parse DOCX {path.name}: {exc}") from exc

        parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return self.assemble(
            filename=path.name,
            mime_type=self.mime_types[0],
            page_texts=["\n".join(parts)],
            metadata={"extractor": "python-docx"},
        )
